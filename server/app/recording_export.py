"""录音报告导出：Word (.docx)、PDF（中文优先使用系统字体）。"""

from __future__ import annotations

import io
import json
import logging
import os
from typing import Any

log = logging.getLogger(__name__)


def _parse_summary(summary_json: str) -> dict[str, Any]:
    if not (summary_json or "").strip():
        return {}
    try:
        return json.loads(summary_json)
    except json.JSONDecodeError:
        return {}


def _windir_fonts() -> str:
    wd = os.environ.get("WINDIR")
    if wd:
        return os.path.join(wd, "Fonts")
    if os.name == "nt":
        return r"C:\Windows\Fonts"
    return "/usr/share/fonts/truetype/dejavu"


def _register_reportlab_cn_font() -> str | None:
    """注册 ReportLab 中文字体，返回注册名；失败返回 None。"""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return None
    fonts_dir = _windir_fonts()
    candidates = [
        ("msyh.ttc", 0),
        ("msyhbd.ttc", 0),
        ("simhei.ttf", None),
        ("simsun.ttc", 0),
    ]
    for i, (fname, subidx) in enumerate(candidates):
        reg_name = f"AiWatchCN{i}"
        path = os.path.join(fonts_dir, fname)
        if not os.path.isfile(path):
            continue
        try:
            if fname.endswith(".ttc") and subidx is not None:
                pdfmetrics.registerFont(TTFont(reg_name, path, subfontIndex=subidx))
            else:
                pdfmetrics.registerFont(TTFont(reg_name, path))
            return reg_name
        except Exception as e:
            log.debug("skip font %s: %s", path, e)
            continue
    return None


def _xml_para(text: str) -> str:
    from xml.sax.saxutils import escape

    t = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    t = escape(t, entities={"\"": "&quot;", "'": "&apos;"})
    return t.replace("\n", "<br/>")


def recording_to_docx(row: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    rid = row.get("id")
    cat = (row.get("category") or "other").strip()
    title = doc.add_heading(f"录音工作报告 · #{rid}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run("分类：").bold = True
    p.add_run(f"{cat}\n")
    p.add_run("设备：").bold = True
    p.add_run(f"{row.get('device_name') or '—'}\n")
    created = row.get("created_at")
    if created is not None:
        p.add_run("创建时间：").bold = True
        p.add_run(f"{created}\n")

    doc.add_heading("一、转写全文", level=1)
    doc.add_paragraph(row.get("transcript") or "（无）")

    summary = _parse_summary(row.get("summary_json") or "")
    kind = (summary.get("kind") or "").strip()
    tasks = summary.get("tasks") if isinstance(summary.get("tasks"), list) else []

    if kind == "work_plan" and tasks:
        doc.add_heading(summary.get("title") or "二、工作计划路线图", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "时间节点"
        hdr[1].text = "任务内容"
        hdr[2].text = "执行人"
        for t in tasks[:50]:
            if not isinstance(t, dict):
                continue
            tc = str(t.get("task_content") or "").strip()
            if not tc:
                continue
            cells = table.add_row().cells
            cells[0].text = str(t.get("time_node") or "待定")
            cells[1].text = tc
            cells[2].text = str(t.get("assignee") or "待定")
        if (summary.get("report") or "").strip():
            doc.add_heading("三、综述", level=1)
            doc.add_paragraph(summary["report"])
    elif summary.get("title"):
        doc.add_heading(summary.get("title") or "二、AI 总结", level=1)
        for h in summary.get("highlights") or []:
            if str(h).strip():
                doc.add_paragraph(str(h).strip(), style="List Bullet")
        if (summary.get("report") or "").strip():
            doc.add_paragraph(summary["report"])
    else:
        doc.add_paragraph("（尚未生成 AI 报告，请稍后刷新或检查 DEEPSEEK_API_KEY）")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def recording_to_pdf(row: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )
    styles = getSampleStyleSheet()
    cn = _register_reportlab_cn_font()
    if cn:
        body = ParagraphStyle(
            "BodyCN",
            parent=styles["Normal"],
            fontName=cn,
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )
        h1 = ParagraphStyle("H1CN", parent=styles["Heading1"], fontName=cn, fontSize=16, leading=20, spaceAfter=10)
    else:
        body = styles["Normal"]
        h1 = styles["Heading1"]
        log.warning("PDF: no Chinese font registered; CJK may show as squares. Install reportlab + Windows fonts.")

    story: list = []
    rid = row.get("id")
    cat = (row.get("category") or "other").strip()
    story.append(Paragraph(_xml_para(f"录音工作报告 · #{rid}"), h1))
    meta = f"分类：{cat}　设备：{row.get('device_name') or '—'}"
    if row.get("created_at") is not None:
        meta += f"　创建时间：{row.get('created_at')}"
    story.append(Paragraph(_xml_para(meta), body))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(_xml_para("一、转写全文"), h1))
    story.append(Paragraph(_xml_para(row.get("transcript") or "（无）"), body))
    story.append(Spacer(1, 0.3 * cm))

    summary = _parse_summary(row.get("summary_json") or "")
    kind = (summary.get("kind") or "").strip()
    tasks = summary.get("tasks") if isinstance(summary.get("tasks"), list) else []

    if kind == "work_plan" and tasks:
        story.append(Paragraph(_xml_para(summary.get("title") or "二、工作计划路线图"), h1))

        def Pcell(t: str) -> Paragraph:
            return Paragraph(_xml_para(str(t)), body)

        data = [
            [Pcell("时间节点"), Pcell("任务内容"), Pcell("执行人")],
        ]
        for t in tasks[:40]:
            if not isinstance(t, dict):
                continue
            tc = str(t.get("task_content") or "").strip()
            if not tc:
                continue
            data.append(
                [
                    Pcell(str(t.get("time_node") or "待定")),
                    Pcell(tc),
                    Pcell(str(t.get("assignee") or "待定")),
                ]
            )
        tw = float(doc.width)
        tbl = Table(data, colWidths=[tw * 0.22, tw * 0.53, tw * 0.25])
        font_tbl = cn or "Helvetica"
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8e8e8")),
                    ("FONTNAME", (0, 0), (-1, -1), font_tbl),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(tbl)
        if (summary.get("report") or "").strip():
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph(_xml_para("三、综述"), h1))
            story.append(Paragraph(_xml_para(summary["report"]), body))
    elif summary.get("title"):
        story.append(Paragraph(_xml_para(summary.get("title") or "二、AI 总结"), h1))
        for h in summary.get("highlights") or []:
            if str(h).strip():
                story.append(Paragraph(_xml_para(f"• {h}"), body))
        if (summary.get("report") or "").strip():
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(_xml_para(summary["report"]), body))
    else:
        story.append(Paragraph(_xml_para("（尚未生成 AI 报告）"), body))

    doc.build(story)
    return buf.getvalue()
